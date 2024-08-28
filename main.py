import openpyxl
from docx import Document
from datetime import datetime

# Carrega a planilha Excel localizada no caminho especificado
workbook = openpyxl.load_workbook("data/fornecedores.xlsx")
sheet_page = workbook["Sheet1"]


# Extrair informações da planilha
for linha in sheet_page.iter_rows(min_row=2, values_only=True):
    nome_fornecedor, endereço, cidade, estado, cep, email = linha

    # Criar um novo documento word
    documento = Document()
    # Adiciona um título ao documento
    documento.add_heading("Contrato de Prestação de Serviço", 0)

    # Cria o texto do contrato utilizando as informações do fornecedor
    texto_contrato = f"""
    Este contrato de prestação de serviços é feito entre {nome_fornecedor}, com endereço em {endereço}, {cidade}, {estado}, CEP {cep}, doravante denominado FORNECEDOR, e a empresa CONTRATANTE.

Pelo presente instrumento particular, as partes têm, entre si, justo e acordado o seguinte:

1. OBJETO DO CONTRATO
    O FORNECEDOR compromete-se a fornecer à CONTRATANTE os serviços/material de acordo com as especificações acordadas, respeitando os padrões de qualidade e os prazos estipulados.

2. PRAZO
    Este contrato tem prazo de vigência de 12 (doze) meses, iniciando-se na data de sua assinatura, podendo ser renovado conforme acordo entre as partes.

3. VALOR E FORMA DE PAGAMENTO
    O valor dos serviços prestados será acordado conforme as demandas da CONTRATANTE e a capacidade de entrega do FORNECEDOR. Os pagamentos serão realizados mensalmente, mediante apresentação de nota fiscal.

4. CONFIDENCIALIDADE
    Todas as informações trocadas entre as partes durante a vigência deste contrato serão tratadas como confidenciais.

Para firmeza e como prova de assim haverem justo e contratado, as partes assinam o presente contrato em duas vias de igual teor e forma.

FORNECEDOR: {nome_fornecedor}
E-mail: {email}

CONTRATANTE: TecnoTech
E-mail: tectectec@tectectec.com

São Paulo, {datetime.now().strftime("%d/%m/%Y")}
    """

    # Adiciona o texto do contrato ao documento
    documento.add_paragraph(texto_contrato)

    # Salva o documento em um diretório específico com o nome do fornecedor
    documento.save(f"data/contratos/contrato_{nome_fornecedor}.docx")
